from django.shortcuts import render
import os
import tempfile
from django.http import Http404, FileResponse, JsonResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
import io
from PIL import Image
from docx.shared import Inches
import uuid
class PdftoDocxConverter(APIView):

    def post(self, request, format=None):

        received_data = request.FILES.get('pdf')
        file_type = request.POST.get('type')

        if not received_data:
            return Response({"error": "No pdf file uploaded."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            

            # Save the uploaded PDF to a temporary file
            temp_pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            temp_pdf_file.write(received_data.read())
            temp_pdf_file.close()

            if file_type == "word":
                # Create a unique filename for the converted DOCX file
                converted_filename = f"{uuid.uuid4().hex}.docx"

                # Define the output file path (you may adjust this)
                output_path = os.path.join("files", converted_filename)

                # Initialize a DOCX document
                doc = Document()

                # Open the PDF file
                pdf_document = fitz.open(temp_pdf_file.name)

                # Create a temporary directory to store images
                temp_image_dir = "temp_images"
                os.makedirs(temp_image_dir, exist_ok=True)

                # Iterate through pages
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)

                    # Create a new page in the DOCX document
                    doc.add_page_break()
                    section = doc.sections[-1]

                    # Set page size (adjust as needed)
                    section.page_width = Pt(612)
                    section.page_height = Pt(792)

                    # Extract text and add it to the DOCX document
                    page_text = page.get_text("text")
                    doc.add_paragraph(page_text)

                    # Extract images and add them to the DOCX document
                    image_list = page.get_images()
                    for img_index, img_matrix in enumerate(image_list):
                        xref = img_matrix[0]
                        img = fitz.Pixmap(pdf_document, xref)
                        
                        img_bytes = img.tobytes("jpeg")
                        img_stream = io.BytesIO(img_bytes)
                        img_path = os.path.join(temp_image_dir, f'temp_image_{page_num}_{img_index}.jpeg')

                        with open(img_path, 'wb') as img_file:
                            img_file.write(img_stream.read())




                        page_width_in_inches = Inches(6.12) 
                        image_width = page_width_in_inches  

                        # Add the image to the DOCX document
                        doc.add_picture(img_path, width=image_width)
                        os.remove(img_path)

                # Remove the temporary image directory
                os.rmdir(temp_image_dir)

                # Save the DOCX file
                doc.save(output_path)

                # Remove the temporary PDF file
                os.remove(temp_pdf_file.name)
               
                #return Response({"converted_file": output_path}, status=status.HTTP_200_OK)
                return JsonResponse({"path": converted_filename})

        except Exception as e:
            print(e)
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
